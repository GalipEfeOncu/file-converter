"""
tests/test_converter.py — FileConverter birim testleri

Issue #10 kapsaminda DOCX -> PDF ve gorsel donusum akislari icin
basari ve hata senaryolarini dogrular.
"""

from pathlib import Path
import sys
from unittest.mock import MagicMock

from docx import Document
from PIL import Image


if "docx2pdf" not in sys.modules:
    mock_docx2pdf = MagicMock()
    mock_docx2pdf.convert = MagicMock()
    sys.modules["docx2pdf"] = mock_docx2pdf

if "pypandoc" not in sys.modules:
    mock_pypandoc = MagicMock()
    mock_pypandoc.convert_file = MagicMock()
    sys.modules["pypandoc"] = mock_pypandoc

from core import converter as converter_module

FileConverter = converter_module.FileConverter


def test_convert_image_rgba_png_to_jpg_success(tmp_path: Path):
    """RGBA PNG girdisi JPG cikisina basariyla donusmeli."""
    source = tmp_path / "input.png"
    output = tmp_path / "output.jpg"

    Image.new("RGBA", (32, 32), color=(255, 0, 0, 128)).save(source, format="PNG")

    result = FileConverter().convert_image(str(source), str(output), "jpg", quality=85)

    assert result is True
    assert output.exists()

    with Image.open(output) as converted:
        assert converted.format == "JPEG"
        assert converted.mode == "RGB"


def test_convert_image_missing_file_returns_false(tmp_path: Path):
    """Olmayan girdi gorseli icin False donmeli."""
    result = FileConverter().convert_image(
        str(tmp_path / "missing.png"),
        str(tmp_path / "output.webp"),
        "webp",
    )

    assert result is False


def test_convert_csv_to_xlsx_success(tmp_path: Path):
    """CSV girdisi XLSX cikisina basariyla donusmeli."""
    source = tmp_path / "input.csv"
    output = tmp_path / "output.xlsx"
    source.write_text("name,value\nalpha,1\nbeta,2\n", encoding="utf-8")

    result = FileConverter().convert_csv_to_xlsx(str(source), str(output))

    assert result is True
    assert output.exists()


def test_convert_csv_to_xlsx_empty_file_returns_false(tmp_path: Path):
    """Bos CSV dosyasi EmptyDataError uzerinden False donmeli."""
    source = tmp_path / "empty.csv"
    output = tmp_path / "output.xlsx"
    source.write_text("", encoding="utf-8")

    result = FileConverter().convert_csv_to_xlsx(str(source), str(output))

    assert result is False
    assert not output.exists()


def test_convert_xlsx_to_csv_success(tmp_path: Path):
    """XLSX girdisi CSV cikisina basariyla donusmeli."""
    source = tmp_path / "input.xlsx"
    output = tmp_path / "output.csv"
    import pandas as pd
    pd.DataFrame([{"col_a": 1, "col_b": 2}]).to_excel(source, index=False, engine="openpyxl")

    result = FileConverter().convert_xlsx_to_csv(str(source), str(output))

    assert result is True
    assert output.exists()
    assert "col_a,col_b" in output.read_text(encoding="utf-8")


def test_convert_xlsx_to_csv_missing_file_returns_false(tmp_path: Path):
    """Olmayan XLSX girdisi icin False donmeli."""
    result = FileConverter().convert_xlsx_to_csv(
        str(tmp_path / "missing.xlsx"),
        str(tmp_path / "output.csv"),
    )

    assert result is False


def test_convert_pdf_to_docx_success(tmp_path: Path, monkeypatch):
    """pdf2docx donusumu basariliysa True donmeli."""
    source = tmp_path / "input.pdf"
    output = tmp_path / "output.docx"
    source.write_bytes(b"%PDF-1.4 fake pdf")
    call_log = {"closed": False}

    class FakeConverter:
        def __init__(self, input_path: str):
            assert input_path == str(source)

        def convert(self, output_path: str, start: int = 0, end=None) -> None:
            assert output_path == str(output)
            assert start == 0
            assert end is None
            Path(output_path).write_bytes(b"fake docx bytes")

        def close(self) -> None:
            call_log["closed"] = True

    monkeypatch.setattr(converter_module, "Converter", FakeConverter)

    result = FileConverter().convert_pdf_to_docx(str(source), str(output))

    assert result is True
    assert output.exists()
    assert call_log["closed"] is True


def test_convert_pdf_to_docx_failure_returns_false(tmp_path: Path, monkeypatch):
    """pdf2docx tarafindaki hata exception yaymadan False donmeli."""
    source = tmp_path / "input.pdf"
    output = tmp_path / "output.docx"
    source.write_bytes(b"%PDF-1.4 fake pdf")

    class FakeConverter:
        def __init__(self, _input_path: str):
            pass

        def convert(self, _output_path: str, start: int = 0, end=None) -> None:
            raise RuntimeError("pdf2docx failure")

        def close(self) -> None:
            return None

    monkeypatch.setattr(converter_module, "Converter", FakeConverter)

    result = FileConverter().convert_pdf_to_docx(str(source), str(output))

    assert result is False
    assert not output.exists()


def test_convert_pdf_to_docx_with_page_range_success(tmp_path: Path, monkeypatch):
    """PDF to DOCX page range parametreleri cagrisi dogru iletmeli."""
    source = tmp_path / "input.pdf"
    output = tmp_path / "output.docx"
    source.write_bytes(b"%PDF-1.4 fake pdf")
    call_log = {"start": None, "end": None}

    class FakeConverter:
        def __init__(self, input_path: str):
            assert input_path == str(source)

        def convert(self, output_path: str, start: int = 0, end=None) -> None:
            call_log["start"] = start
            call_log["end"] = end
            Path(output_path).write_bytes(b"fake docx bytes")

        def close(self) -> None:
            return None

    monkeypatch.setattr(converter_module, "Converter", FakeConverter)

    result = FileConverter().convert_pdf_to_docx(str(source), str(output), start=1, end=2)

    assert result is True
    assert output.exists()
    assert call_log["start"] == 1
    assert call_log["end"] == 2


def test_convert_image_quality_preset_string_success(tmp_path: Path):
    """Quality preset string ile gorsel donusumu basarili olmali."""
    source = tmp_path / "input.png"
    output = tmp_path / "output.jpg"

    Image.new("RGB", (32, 32), color=(255, 0, 0)).save(source, format="PNG")

    result = FileConverter().convert_image(str(source), str(output), "jpg", quality="high")

    assert result is True
    assert output.exists()


def test_convert_image_invalid_quality_preset_returns_false(tmp_path: Path):
    """Gecersiz quality preset'i False donmeli."""
    source = tmp_path / "input.png"
    output = tmp_path / "output.jpg"
    Image.new("RGB", (32, 32), color=(255, 255, 255)).save(source, format="PNG")

    result = FileConverter().convert_image(str(source), str(output), "jpg", quality="super")

    assert result is False
    assert not output.exists()


def test_convert_rtf_to_docx_success(tmp_path: Path, monkeypatch):
    """RTF to DOCX donusumu pypandoc ile basarili olmali."""
    source = tmp_path / "input.rtf"
    output = tmp_path / "output.docx"
    source.write_text("{\\rtf1\\ansi Hello}", encoding="utf-8")

    def fake_convert_file(input_path: str, to: str, format: str, outputfile: str):
        assert input_path == str(source)
        assert to == "docx"
        assert format == "rtf"
        assert outputfile == str(output)
        Path(outputfile).write_bytes(b"fake docx bytes")

    monkeypatch.setattr(converter_module.pypandoc, "convert_file", fake_convert_file)
    result = FileConverter().convert_rtf_to_docx(str(source), str(output))

    assert result is True
    assert output.exists()


def test_convert_odt_to_docx_success(tmp_path: Path, monkeypatch):
    """ODT to DOCX donusumu pypandoc ile basarili olmali."""
    source = tmp_path / "input.odt"
    output = tmp_path / "output.docx"
    source.write_text("dummy odt content", encoding="utf-8")

    def fake_convert_file(input_path: str, to: str, format: str, outputfile: str):
        assert input_path == str(source)
        assert to == "docx"
        assert format == "odt"
        assert outputfile == str(output)
        Path(outputfile).write_bytes(b"fake docx bytes")

    monkeypatch.setattr(converter_module.pypandoc, "convert_file", fake_convert_file)
    result = FileConverter().convert_odt_to_docx(str(source), str(output))

    assert result is True
    assert output.exists()


def test_convert_docx_to_pdf_success(tmp_path: Path, monkeypatch):
    """docx2pdf cagrisi basariliysa True donmeli ve cikti dosyasi olusmali."""
    source = tmp_path / "input.docx"
    output = tmp_path / "output.pdf"

    document = Document()
    document.add_paragraph("Donusturme testi")
    document.save(source)

    def fake_convert(input_path: str, output_path: str) -> None:
        assert input_path == str(source)
        Path(output_path).write_bytes(b"%PDF-1.4 fake pdf")

    monkeypatch.setattr(converter_module, "docx2pdf_convert", fake_convert)

    result = FileConverter().convert_docx_to_pdf(str(source), str(output))

    assert result is True
    assert output.exists()
    assert output.read_bytes().startswith(b"%PDF")


def test_convert_docx_to_pdf_failure_returns_false(tmp_path: Path, monkeypatch):
    """docx2pdf hata verirse metot exception yaymadan False donmeli."""
    source = tmp_path / "input.docx"
    output = tmp_path / "output.pdf"

    document = Document()
    document.add_paragraph("Basarisiz donusum testi")
    document.save(source)

    def fake_convert(_input_path: str, _output_path: str) -> None:
        raise RuntimeError("MS Word bulunamadi")

    monkeypatch.setattr(converter_module, "docx2pdf_convert", fake_convert)

    result = FileConverter().convert_docx_to_pdf(str(source), str(output))

    assert result is False
    assert not output.exists()


def test_convert_docx_to_txt_success(tmp_path: Path):
    """DOCX icerigi UTF-8 TXT cikisina yazilmali."""
    source = tmp_path / "input.docx"
    output = tmp_path / "output.txt"

    document = Document()
    document.add_paragraph("Ilk satir")
    document.add_paragraph("")
    document.add_paragraph("Ikinci satir")
    document.save(source)

    result = FileConverter().convert_docx_to_txt(str(source), str(output))

    assert result is True
    assert output.exists()
    assert output.read_text(encoding="utf-8") == "Ilk satir\nIkinci satir\n"


def test_convert_docx_to_txt_missing_file_returns_false(tmp_path: Path):
    """Olmayan DOCX girdisi icin False donmeli."""
    result = FileConverter().convert_docx_to_txt(
        str(tmp_path / "missing.docx"),
        str(tmp_path / "output.txt"),
    )

    assert result is False
def test_pdf_to_images_fail(): 
    converter = FileConverter()
    assert converter.pdf_to_images("yok.pdf", "out") == []

def test_pdf_to_images_success(tmp_path: Path, monkeypatch):
    """PDF to images donusumu basariliysa kaydedilen dosya yollarini donmeli."""
    source = tmp_path / "input.pdf"
    output_dir = tmp_path / "images"
    source.write_bytes(b"%PDF-1.4 fake pdf")
    saved_files = []

    class FakePage:
        def get_pixmap(self, matrix=None):
            class FakePixmap:
                def save(self, path):
                    saved_files.append(path)
            return FakePixmap()

    class FakeDoc:
        def __init__(self, path):
            self.pages = [FakePage(), FakePage()]  # 2 sayfa

        def __iter__(self):
            return iter(self.pages)

        def close(self):
            pass

    monkeypatch.setattr(converter_module, "fitz", MagicMock())
    monkeypatch.setattr(converter_module.fitz, "open", FakeDoc)
    monkeypatch.setattr(converter_module.fitz, "Matrix", lambda x, y: None)

    result = FileConverter().pdf_to_images(str(source), str(output_dir))

    assert len(result) == 2
    assert "p_1.png" in result[0]
    assert "p_2.png" in result[1]


def test_merge_pdfs_fail():
    converter = FileConverter()
    assert converter.merge_pdfs([], "out.pdf") is False

def test_merge_pdfs_success(tmp_path: Path, monkeypatch):
    """PDF merge basariliysa True donmeli."""
    input1 = tmp_path / "input1.pdf"
    input2 = tmp_path / "input2.pdf"
    output = tmp_path / "merged.pdf"
    input1.write_bytes(b"%PDF-1.4 fake1")
    input2.write_bytes(b"%PDF-1.4 fake2")

    class FakeInputDoc:
        def __enter__(self):
            return self
        def __exit__(self, *args):
            pass

    class FakeOutputDoc:
        def __init__(self):
            self.inserted = []

        def insert_pdf(self, doc):
            self.inserted.append(doc)

        def save(self, path):
            Path(path).write_bytes(b"%PDF-1.4 merged")

        def close(self):
            pass

    def fake_open(path=None):
        if path is None:
            return FakeOutputDoc()
        else:
            return FakeInputDoc()

    monkeypatch.setattr(converter_module, "fitz", MagicMock())
    monkeypatch.setattr(converter_module.fitz, "open", fake_open)

    result = FileConverter().merge_pdfs([str(input1), str(input2)], str(output))

    assert result is True
    assert output.exists()


def test_batch_convert_success(tmp_path: Path, monkeypatch):
    """Batch convert basariliysa her dosya icin True donmeli."""
    input1 = tmp_path / "input1.png"
    input2 = tmp_path / "input2.png"
    output_dir = tmp_path / "output"
    input1.write_bytes(b"fake png1")
    input2.write_bytes(b"fake png2")

    # convert_image'i mock et
    def fake_convert_image(self, input_path, output_path, target_format, **kwargs):
        Path(output_path).write_bytes(b"converted")
        return True

    monkeypatch.setattr(FileConverter, "convert_image", fake_convert_image)

    result = FileConverter().batch_convert([str(input1), str(input2)], str(output_dir), "jpg")

    assert result[str(input1)] is True
    assert result[str(input2)] is True
    assert len(result) == 2


def test_batch_convert_failure(tmp_path: Path, monkeypatch):
    """Batch convert hata verirse False donmeli."""
    input1 = tmp_path / "input1.png"
    output_dir = tmp_path / "output"
    input1.write_bytes(b"fake png")

    # convert_image'i mock et - hata versin
    def fake_convert_image(self, input_path, output_path, target_format, **kwargs):
        return False

    monkeypatch.setattr(FileConverter, "convert_image", fake_convert_image)

    result = FileConverter().batch_convert([str(input1)], str(output_dir), "jpg")

    assert result[str(input1)] is False
    assert len(result) == 1


def test_convert_rtf_to_docx_failure(tmp_path: Path, monkeypatch):
    """RTF to DOCX donusumu hata verirse False donmeli."""
    source = tmp_path / "input.rtf"
    output = tmp_path / "output.docx"
    source.write_text("{\\rtf1\\ansi Hello}", encoding="utf-8")

    def fake_convert_file(input_path: str, to: str, format: str, outputfile: str):
        raise RuntimeError("pypandoc error")

    monkeypatch.setattr(converter_module.pypandoc, "convert_file", fake_convert_file)
    result = FileConverter().convert_rtf_to_docx(str(source), str(output))

    assert result is False
    assert not output.exists()


def test_convert_odt_to_docx_failure(tmp_path: Path, monkeypatch):
    """ODT to DOCX donusumu hata verirse False donmeli."""
    source = tmp_path / "input.odt"
    output = tmp_path / "output.docx"
    source.write_text("dummy odt content", encoding="utf-8")

    def fake_convert_file(input_path: str, to: str, format: str, outputfile: str):
        raise RuntimeError("pypandoc error")

    monkeypatch.setattr(converter_module.pypandoc, "convert_file", fake_convert_file)
    result = FileConverter().convert_odt_to_docx(str(source), str(output))

    assert result is False
    assert not output.exists()