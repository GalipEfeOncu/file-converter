"""
tests/test_viewer.py — FileViewer display_* metotları ve _dispatch_viewer testleri
Issue #13 + Issue #8 + Issue #24 + Issue #29 — Abdulkadir Sar (Aksar712)

Kapsam:
  - display_image başarılı senaryo (st.image çağrısını doğrular)
  - display_image eksik dosya senaryosu (exception fırlatmaz)
  - display_text_document: .txt, .py (kod), .docx başarı ve hata senaryoları
  - _dispatch_viewer: görsel, ses, metin ve desteklenmeyen uzantı yönlendirmeleri
  - render_pdf pagination testleri (Issue #24)
  - display_table arama/filtre testleri (Issue #24)
  - Issue #29: cache, zoom, metadata, extract_text testleri
"""

import sys
import os
import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch

# ---------------------------------------------------------------------------
# PyMuPDF (fitz) yoksa modül seviyesinde stub ekle — CI/lokal uyumluluğu için
# ---------------------------------------------------------------------------
if "fitz" not in sys.modules:
    fitz_stub = MagicMock()
    sys.modules["fitz"] = fitz_stub

from docx import Document
from core.viewer import FileViewer
from core import viewer as viewer_module

# ---------------------------------------------------------------------------
# Yardımcı: tek piksellik minimum PNG bayt dizisi
# ---------------------------------------------------------------------------
MINIMAL_PNG = (
    b"\x89PNG\r\n\x1a\n"
    b"\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx"
    b"\x9cc\xf8\x0f\x00\x00\x01\x01\x00\x05\x18\xd8N\x00"
    b"\x00\x00\x00IEND\xaeB`\x82"
)


# ===========================================================================
# display_image testleri
# ===========================================================================

class TestDisplayImage:
    """FileViewer.display_image metodu için testler."""

    def test_display_image_calls_st_image_success(self, tmp_path: Path):
        """Başarılı senaryo: gerçek PNG dosyası ile st.image çağrısını doğrular."""
        from core.viewer import FileViewer

        img_file = tmp_path / "test.png"
        img_file.write_bytes(MINIMAL_PNG)

        fv = FileViewer()
        with patch("core.viewer.st.image") as mock_st_image:
            # st.radio'yu mockla ki "Fit" dönsün (varsayılan)
            with patch("core.viewer.st.radio", return_value="Fit"):
                fv.display_image(str(img_file))
                mock_st_image.assert_called_once_with(str(img_file), use_container_width=True)

    def test_display_image_missing_file_does_not_raise(self, tmp_path: Path):
        """Hata senaryosu: eksik dosya exception fırlatmamalı; st.error çağrılmalı."""
        from core.viewer import FileViewer

        missing = str(tmp_path / "nonexistent.png")
        fv = FileViewer()

        with (
            patch("core.viewer.st.image", side_effect=FileNotFoundError("Dosya yok")),
            patch("core.viewer.st.error") as mock_error,
        ):
            # Exception fırlatmamalı
            fv.display_image(missing)
            assert mock_error.called, "Hata durumunda st.error çağrılmalıydı."


# ===========================================================================
# _dispatch_viewer testleri (Dashboard üzerinde)
# ===========================================================================

class TestDispatchViewer:
    """Dashboard._dispatch_viewer için testler."""

    def _make_uploaded_file(self, name: str, content: bytes = b"dummy"):
        """Sahte bir Streamlit UploadedFile nesnesi üretir."""
        mock_file = MagicMock()
        mock_file.name = name
        mock_file.getbuffer.return_value = content
        return mock_file

    def test_dispatch_viewer_routes_image_to_display_image(self, tmp_path: Path):
        """Başarılı senaryo: .png uzantısı display_image'e yönlendirilmeli."""
        from ui.dashboard import Dashboard

        texts = {}
        dash = Dashboard(texts)

        uploaded = self._make_uploaded_file("photo.png", MINIMAL_PNG)
        fake_path = str(tmp_path / "photo.png")

        mock_fv = MagicMock()
        mock_fv_class = MagicMock(return_value=mock_fv)

        # FileViewer modül seviyesinde import edildiği için ui.dashboard.FileViewer
        # ile patch edilebilir
        with (
            patch.object(Dashboard, "_save_upload_to_temp", return_value=fake_path),
            patch("ui.dashboard.FileViewer", mock_fv_class),
            patch("pathlib.Path.unlink"),
        ):
            dash._dispatch_viewer(uploaded)

        mock_fv.display_image.assert_called_once_with(fake_path)


    def test_dispatch_viewer_unsupported_ext_shows_warning(self, tmp_path: Path):
        """Hata senaryosu: desteklenmeyen uzantı için st.warning çağrılmalı."""
        from ui.dashboard import Dashboard

        texts = {"error_unsupported_file": "Desteklenmeyen dosya türü."}
        dash = Dashboard(texts)

        uploaded = self._make_uploaded_file("archive.zip", b"PK")
        fake_path = str(tmp_path / "archive.zip")

        mock_fv = MagicMock()
        mock_fv_class = MagicMock(return_value=mock_fv)

        with (
            patch.object(Dashboard, "_save_upload_to_temp", return_value=fake_path),
            patch("ui.dashboard.FileViewer", mock_fv_class),
            patch("streamlit.warning") as mock_warning,
            patch("pathlib.Path.unlink"),
        ):
            dash._dispatch_viewer(uploaded)

        mock_warning.assert_called_once()
        # Uyarı mesajı i18n'den gelmeli
        assert "Desteklenmeyen" in mock_warning.call_args[0][0]

    def test_dispatch_viewer_routes_audio_to_display_audio(self, tmp_path: Path):
        """Başarılı senaryo: .mp3 uzantısı display_audio'ya yönlendirilmeli."""
        from ui.dashboard import Dashboard

        texts = {}
        dash = Dashboard(texts)
        uploaded = self._make_uploaded_file("song.mp3", b"\xff\xfb")
        fake_path = str(tmp_path / "song.mp3")

        mock_fv = MagicMock()
        mock_fv_class = MagicMock(return_value=mock_fv)

        with (
            patch.object(Dashboard, "_save_upload_to_temp", return_value=fake_path),
            patch("ui.dashboard.FileViewer", mock_fv_class),
            patch("pathlib.Path.unlink"),
        ):
            dash._dispatch_viewer(uploaded)

        assert mock_fv.display_audio.called, "display_audio çağrılmalıydı."

    def test_dispatch_viewer_routes_docx_to_display_text(self, tmp_path: Path):
        """Başarılı senaryo: .docx uzantısı display_text_document'a yönlendirilmeli."""
        from ui.dashboard import Dashboard

        texts = {}
        dash = Dashboard(texts)
        uploaded = self._make_uploaded_file("report.docx", b"PK")
        fake_path = str(tmp_path / "report.docx")

        mock_fv = MagicMock()
        mock_fv_class = MagicMock(return_value=mock_fv)

        with (
            patch.object(Dashboard, "_save_upload_to_temp", return_value=fake_path),
            patch("ui.dashboard.FileViewer", mock_fv_class),
            patch("pathlib.Path.unlink"),
            patch("streamlit.spinner", return_value=MagicMock(__enter__=lambda s: s, __exit__=MagicMock(return_value=False))),
        ):
            dash._dispatch_viewer(uploaded)

        assert mock_fv.display_text_document.called, "display_text_document çağrılmalıydı."


# ===========================================================================
# display_text_document testleri (Issue #8)
# ===========================================================================

class TestDisplayTextDocument:
    """FileViewer.display_text_document genişletilmiş davranışı için testler."""

    def test_display_txt_file_success(self, tmp_path: Path):
        """Başarılı senaryo: .txt dosyası st.text_area ile gösterilmeli."""
        from core.viewer import FileViewer

        txt_file = tmp_path / "notes.txt"
        txt_file.write_text("Merhaba Dünya", encoding="utf-8")

        fv = FileViewer()
        with patch("core.viewer.st.text_area") as mock_area:
            fv.display_text_document(str(txt_file))
            assert mock_area.called, "st.text_area çağrılmalıydı."
            # İçerik doğru geçirilmeli
            assert "Merhaba Dünya" in mock_area.call_args[0][1]

    def test_display_python_file_uses_st_code(self, tmp_path: Path):
        """Başarılı senaryo: .py dosyası st.code(language='python') ile gösterilmeli."""
        from core.viewer import FileViewer

        py_file = tmp_path / "script.py"
        py_file.write_text("print('hello')", encoding="utf-8")

        fv = FileViewer()
        with patch("core.viewer.st.code") as mock_code:
            fv.display_text_document(str(py_file))
            mock_code.assert_called_once_with("print('hello')", language="python")

    def test_display_text_document_read_error_does_not_raise(self, tmp_path: Path):
        """Hata senaryosu: okuma hatası exception fırlatmamalı; st.error çağrılmalı."""
        from core.viewer import FileViewer

        py_file = tmp_path / "broken.py"
        py_file.write_text("x = 1", encoding="utf-8")

        fv = FileViewer()
        with (
            patch("builtins.open", side_effect=PermissionError("Erişim reddedildi")),
            patch("core.viewer.st.error") as mock_error,
        ):
            fv.display_text_document(str(py_file))
            assert mock_error.called, "Hata durumunda st.error çağrılmalıydı."

    def test_extract_text_from_txt(self, tmp_path: Path):
        """Başarılı senaryo: extract_text ile .txt dosyasından metin çıkarma."""
        from core.viewer import FileViewer

        txt_file = tmp_path / "sample.txt"
        txt_file.write_text("Örnek Metin İçeriği", encoding="utf-8")

        fv = FileViewer()
        extracted = fv.extract_text(str(txt_file))
        assert extracted == "Örnek Metin İçeriği", "Çıkarılan metin doğru olmalı."

    def test_extract_text_unsupported_ext(self, tmp_path: Path):
        """Hata senaryosu: desteklenmeyen uzantı için boş string döner."""
        from core.viewer import FileViewer

        unsupported_file = tmp_path / "image.png"
        unsupported_file.write_bytes(MINIMAL_PNG)

        fv = FileViewer()
        extracted = fv.extract_text(str(unsupported_file))
        assert extracted == "", "Desteklenmeyen dosya türünden metin çıkarılamamalı."


# ===========================================================================
# Issue #24: render_pdf parametreleri + display_table arama testleri
# ===========================================================================

class TestRenderPdfPagination:
    """render_pdf start/end parametre desteği testleri (Issue #24)."""

    def test_render_pdf_start_end_limits_pages(self, tmp_path: Path):
        """Başarılı senaryo: render_pdf start=0, end=1 ile yalnızca 1 sayfa döner."""
        from core.viewer import FileViewer

        fv = FileViewer()
        # Gerçek bir PyMuPDF belgesi yerine mock kullan
        mock_doc = MagicMock()
        mock_page = MagicMock()
        mock_pixmap = MagicMock()
        mock_pixmap.tobytes.return_value = b"fake_png_bytes"
        mock_page.get_pixmap.return_value = mock_pixmap
        mock_doc.load_page.return_value = mock_page
        mock_doc.__len__ = MagicMock(return_value=10)

        from core.viewer import st
        st.cache_data.clear()

        with patch("core.viewer.fitz.open", return_value=mock_doc):
            result = fv.render_pdf("fake.pdf", start=0, end=1)

        assert len(result) == 1, "start=0, end=1 ile yalnızca 1 sayfa dönmeli."
        assert result[0] == b"fake_png_bytes"

    def test_render_pdf_default_returns_all_pages(self, tmp_path: Path):
        """Başarılı senaryo: parametresiz çağrıda tüm sayfalar döner."""
        from core.viewer import FileViewer

        fv = FileViewer()
        PAGE_COUNT = 5
        mock_doc = MagicMock()
        mock_page = MagicMock()
        mock_pixmap = MagicMock()
        mock_pixmap.tobytes.return_value = b"page"
        mock_page.get_pixmap.return_value = mock_pixmap
        mock_doc.load_page.return_value = mock_page
        mock_doc.__len__ = MagicMock(return_value=PAGE_COUNT)

        from core.viewer import st
        st.cache_data.clear()

        with patch("core.viewer.fitz.open", return_value=mock_doc):
            result = fv.render_pdf("fake.pdf")

        assert len(result) == PAGE_COUNT, "Varsayılan çağrıda tüm sayfalar dönmeli."

    def test_render_pdf_out_of_bound_end_clamps(self, tmp_path: Path):
        """Sınır senaryosu: end > toplam sayfa olursa son sayfada kesilir."""
        from core.viewer import FileViewer

        fv = FileViewer()
        mock_doc = MagicMock()
        mock_page = MagicMock()
        mock_pixmap = MagicMock()
        mock_pixmap.tobytes.return_value = b"p"
        mock_page.get_pixmap.return_value = mock_pixmap
        mock_doc.load_page.return_value = mock_page
        mock_doc.__len__ = MagicMock(return_value=3)

        from core.viewer import st
        st.cache_data.clear()

        with patch("core.viewer.fitz.open", return_value=mock_doc):
            result = fv.render_pdf("fake.pdf", start=0, end=999)

        assert len(result) == 3, "end > toplam sayfa olunca toplam sayfa kadar dönmeli."


class TestDisplayTableSearch:
    """display_table arama/filtre özelliği testleri (Issue #24)."""

    def test_display_table_no_query_shows_full_dataframe(self, tmp_path: Path):
        """Başarılı senaryo: arama sorgusu yokken tüm DataFrame gösterilmeli."""
        from core.viewer import FileViewer
        import pandas as pd

        csv_file = tmp_path / "data.csv"
        csv_file.write_text("isim,yas\nAli,25\nVeli,30\n", encoding="utf-8")

        fv = FileViewer()
        with (
            patch("core.viewer.st.text_input", return_value=""),  # boş sorgu
            patch("core.viewer.st.dataframe") as mock_df,
            patch("core.viewer.st.caption"),
        ):
            fv.display_table(str(csv_file))
            assert mock_df.called, "Sorgu yokken st.dataframe çağrılmalı."
            called_df = mock_df.call_args[0][0]
            assert len(called_df) == 2, "Tüm satırlar görüntülenmeli."

    def test_display_table_query_filters_rows(self, tmp_path: Path):
        """Başarılı senaryo: arama sorgusu eşleşen satırları filtreler."""
        from core.viewer import FileViewer

        csv_file = tmp_path / "data.csv"
        csv_file.write_text("isim,yas\nAli,25\nVeli,30\n", encoding="utf-8")

        fv = FileViewer()
        with (
            patch("core.viewer.st.text_input", return_value="Ali"),
            patch("core.viewer.st.dataframe") as mock_df,
            patch("core.viewer.st.caption"),
        ):
            fv.display_table(str(csv_file))
            assert mock_df.called, "Eşleşme varsa st.dataframe çağrılmalı."
            called_df = mock_df.call_args[0][0]
            assert len(called_df) == 1, "Yalnızca 'Ali' satırı dönmeli."

    def test_display_table_no_match_shows_info(self, tmp_path: Path):
        """Hata senaryosu: eşleşme yoksa st.info çağrılmalı."""
        from core.viewer import FileViewer

        csv_file = tmp_path / "data.csv"
        csv_file.write_text("isim,yas\nAli,25\nVeli,30\n", encoding="utf-8")

        texts = {"table_no_match": "Aramanızla eşleşen satır bulunamadı."}
        fv = FileViewer()
        with (
            patch("core.viewer.st.text_input", return_value="XYZNOMATCH"),
            patch("core.viewer.st.info") as mock_info,
            patch("core.viewer.st.caption"),
        ):
            fv.display_table(str(csv_file), texts=texts)
            assert mock_info.called, "Eşleşme yokken st.info çağrılmalı."
            assert "eşleşen" in mock_info.call_args[0][0], "i18n mesajı gösterilmeli."


# ===========================================================================
# Issue #29 — Cache, Zoom, Metadata testleri
# ===========================================================================


def test_cached_render_pdf_exists_as_module_level_function():
    """Cache'li render fonksiyonu modul seviyesinde tanimli olmali."""
    import core.viewer as vm
    assert callable(getattr(vm, "_cached_render_pdf", None)), "_cached_render_pdf bulunamadi"
    assert callable(getattr(vm, "_cached_read_table", None)), "_cached_read_table bulunamadi"


def test_display_image_fit_mode(monkeypatch, tmp_path: Path):
    """display_image Fit modunda use_container_width=True ile st.image cagirilmali."""
    img_path = tmp_path / "sample.png"
    img_path.write_bytes(b"fake-png-bytes")
    calls = {}

    monkeypatch.setattr(
        viewer_module.st,
        "radio",
        lambda label, options, **kwargs: "Fit"
    )
    monkeypatch.setattr(
        viewer_module.st,
        "image",
        lambda data, **kwargs: calls.update({"data": data, "kwargs": kwargs})
    )

    FileViewer().display_image(str(img_path))

    assert calls["kwargs"].get("use_container_width") is True
    assert calls["data"] == str(img_path)


def test_display_image_100_mode(monkeypatch, tmp_path: Path):
    """display_image 100% modunda use_container_width=False ile st.image cagirilmali."""
    img_path = tmp_path / "sample.png"
    img_path.write_bytes(b"fake-png-bytes")
    calls = {}

    monkeypatch.setattr(
        viewer_module.st,
        "radio",
        lambda label, options, **kwargs: "100%"
    )
    monkeypatch.setattr(
        viewer_module.st,
        "image",
        lambda data, **kwargs: calls.update({"data": data, "kwargs": kwargs})
    )

    FileViewer().display_image(str(img_path))

    assert calls["kwargs"].get("use_container_width") is False


def test_display_image_200_mode(monkeypatch, tmp_path: Path):
    """display_image 200% modunda width=1200 ile st.image cagirilmali."""
    img_path = tmp_path / "sample.png"
    img_path.write_bytes(b"fake-png-bytes")
    calls = {}

    monkeypatch.setattr(
        viewer_module.st,
        "radio",
        lambda label, options, **kwargs: "200%"
    )
    monkeypatch.setattr(
        viewer_module.st,
        "image",
        lambda data, **kwargs: calls.update({"data": data, "kwargs": kwargs})
    )

    FileViewer().display_image(str(img_path))

    assert calls["kwargs"].get("width") == 1200


def test_extract_text_from_txt_module(tmp_path: Path):
    """extract_text TXT dosyasindan icerik dogrudan donmeli."""
    txt_path = tmp_path / "sample.txt"
    txt_path.write_text("Merhaba dunya\nIkinci satir", encoding="utf-8")

    result = FileViewer().extract_text(str(txt_path))

    assert "Merhaba dunya" in result
    assert "Ikinci satir" in result


def test_extract_text_from_docx(tmp_path: Path):
    """extract_text DOCX dosyasindan paragraf metinleri birlestirilmis donmeli."""
    docx_path = tmp_path / "sample.docx"
    belge = Document()
    belge.add_paragraph("Birinci paragraf")
    belge.add_paragraph("Ikinci paragraf")
    belge.save(docx_path)

    result = FileViewer().extract_text(str(docx_path))

    assert "Birinci paragraf" in result
    assert "Ikinci paragraf" in result


def test_extract_text_unsupported_returns_empty(tmp_path: Path):
    """extract_text desteklenmeyen uzanti icin bos string donmeli."""
    json_path = tmp_path / "data.json"
    json_path.write_text('{"key": "value"}', encoding="utf-8")

    result = FileViewer().extract_text(str(json_path))

    assert result == ""


def test_display_table_shows_metadata(monkeypatch, tmp_path: Path):
    """display_table st.caption ile metadata ozeti gostermeli."""
    csv_path = tmp_path / "meta.csv"
    csv_path.write_text("name,value,score\nalpha,1,3.14\n", encoding="utf-8")
    captions = []

    monkeypatch.setattr(viewer_module.st, "caption", captions.append)
    monkeypatch.setattr(
        viewer_module.st, "dataframe", lambda df, use_container_width=True: None
    )
    monkeypatch.setattr(
        viewer_module.st, "text_input", lambda label, key=None: ""
    )

    FileViewer().display_table(str(csv_path))

    assert captions, "st.caption cagirilmadi"
    caption_text = captions[0]
    assert "satir" in caption_text or "rows" in caption_text or "1" in caption_text
    assert "sutun" in caption_text or "columns" in caption_text or "3" in caption_text
