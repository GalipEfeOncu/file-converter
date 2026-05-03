"""
core/viewer.py — Dosya Görüntüleme ve Önizleme Modülü
Sahibi: Abdulkadir Sar (Görüntüleme Uzmanı)
"""

import logging
import mimetypes
import os

import fitz
import pandas as pd
import streamlit as st
import docx  # python-docx


@st.cache_data(ttl=3600)
def _cached_render_pdf(pdf_path: str) -> list[bytes]:
    """Cache'li PDF render yardımcısı."""
    logging.debug("Cache MISS: PDF render başlatılıyor — %s", pdf_path)
    doc = fitz.open(pdf_path)
    resim_listesi: list[bytes] = []
    for sayfa_numarasi in range(len(doc)):
        sayfa = doc.load_page(sayfa_numarasi)
        resim_verisi = sayfa.get_pixmap()
        png_formati = resim_verisi.tobytes("png")
        resim_listesi.append(png_formati)
    logging.debug("Cache'e yazıldı: %d sayfa — %s", len(resim_listesi), pdf_path)
    return resim_listesi


@st.cache_data(ttl=3600)
def _cached_read_table(file_path: str) -> pd.DataFrame:
    """Cache'li tablo okuyucu."""
    logging.debug("Cache MISS: Tablo okunuyor — %s", file_path)
    _, uzanti = os.path.splitext(file_path)
    uzanti = uzanti.lower()
    if uzanti == ".csv":
        df = pd.read_csv(file_path)
    elif uzanti in [".xls", ".xlsx"]:
        df = pd.read_excel(file_path)
    else:
        raise ValueError("Unsupported format")
    logging.debug("Cache'e yazıldı: %d×%d tablo — %s", len(df), len(df.columns), file_path)
    return df


class FileViewer:
    """Dosyaları önizleme için uygun formata dönüştürür ve arayüzde gösterir."""

    def render_pdf(self, pdf_path: str, start: int = 0, end: int | None = None) -> list[bytes]:
        """PDF sayfalarini gorsel olarak render eder."""
        all_pages = _cached_render_pdf(pdf_path)
        total = len(all_pages)
        finish = end if end is not None else total
        return all_pages[max(0, start):min(finish, total)]

    def read_table(self, file_path: str) -> pd.DataFrame:
        """CSV veya Excel dosyalarını DataFrame olarak okur."""
        try:
            return _cached_read_table(file_path)
        except Exception as e:
            raise ValueError(f"Error reading file: {e}") from e

    def extract_text(self, file_path: str) -> str:
        """Metin tabanlı dosyalardan ham metin çıkarır."""
        dosya_adi, uzanti = os.path.splitext(file_path)
        uzanti = uzanti.lower()
        metin = ""

        try:
            if uzanti == '.pdf':
                doc = fitz.open(file_path)
                for page in doc:
                    metin += page.get_text() + "\n"
            elif uzanti in ['.docx', '.doc']:
                doc = docx.Document(file_path)
                metin = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            elif uzanti in ['.txt', '.csv']:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    metin = f.read()
        except Exception as e:
            logging.error(f"Error extracting text ({file_path}): {e}")

        return metin.strip()

    def display_pdf(self, pdf_path: str, texts: dict | None = None) -> None:
        """PDF sayfalarini Streamlit arayuzunde sayfa sayfa gosterir."""
        if texts is None:
            texts = {}

        try:
            doc = fitz.open(pdf_path)
            toplam_sayfa = len(doc)
        except Exception as e:
            st.error(f"Error: {e}")
            return

        if toplam_sayfa == 0:
            return

        if "pdf_current_page" not in st.session_state:
            st.session_state["pdf_current_page"] = 0

        if st.session_state.get("_pdf_last_path") != pdf_path:
            st.session_state["pdf_current_page"] = 0
            st.session_state["_pdf_last_path"] = pdf_path

        current = st.session_state["pdf_current_page"]
        pdf_page_of = texts.get("pdf_page_of", "of")

        col_info, col_nav = st.columns([3, 2])
        with col_info:
            sayfa_no = st.number_input(
                f"{texts.get('settings_language_label', 'Page')} ({current + 1} {pdf_page_of} {toplam_sayfa})",
                min_value=1,
                max_value=toplam_sayfa,
                value=current + 1,
                step=1,
                key="pdf_page_input",
            )
            current = sayfa_no - 1
            st.session_state["pdf_current_page"] = current

        with col_nav:
            st.write("<div style='height: 28px;'></div>", unsafe_allow_html=True)
            nav_col1, nav_col2 = st.columns(2)
            with nav_col1:
                if st.button(texts.get("pdf_prev", "Previous"), key="pdf_prev", disabled=(current == 0), use_container_width=True):
                    st.session_state["pdf_current_page"] = max(0, current - 1)
                    st.rerun()
            with nav_col2:
                if st.button(texts.get("pdf_next", "Next"), key="pdf_next", disabled=(current >= toplam_sayfa - 1), use_container_width=True):
                    st.session_state["pdf_current_page"] = min(toplam_sayfa - 1, current + 1)
                    st.rerun()

        resimler = self.render_pdf(pdf_path, start=current, end=current + 1)
        if resimler:
            st.image(resimler[0], caption=f"{current + 1} / {toplam_sayfa}", use_container_width=True)

    def display_table(self, file_path: str, texts: dict | None = None) -> None:
        """Tablo verilerini Streamlit arayuzunde gosterir."""
        if texts is None:
            texts = {}

        try:
            df = self.read_table(file_path)
            dtype_str = ", ".join(f"{dtype}" for dtype in df.dtypes.unique())
            st.caption(f"{len(df)} x {len(df.columns)} · {dtype_str}")
        except ValueError as e:
            st.error(str(e))
            return

        search_label = texts.get("pdf_search_label", "Search...")
        query = st.text_input(search_label, key="table_search_query")

        if query:
            filtered = df[df.apply(lambda row: row.astype(str).str.contains(query, case=False, na=False).any(), axis=1)]
            if filtered.empty:
                st.info(texts.get("table_no_match", "No match found."))
            else:
                st.dataframe(filtered, use_container_width=True)
        else:
            st.dataframe(df, use_container_width=True)

    def display_image(self, file_path: str, texts: dict | None = None) -> None:
        """Görsel dosyaları zoom kontrolü ile Streamlit arayüzünde gösterir."""
        if texts is None:
            texts = {}
            
        zoom_options = {
            "Fit": texts.get("image_zoom_fit", "Fit"),
            "100%": "100%",
            "200%": "200%"
        }
        
        selected_zoom_key = st.radio(
            "Zoom",
            options=list(zoom_options.keys()),
            format_func=lambda x: zoom_options[x],
            horizontal=True,
            label_visibility="collapsed",
            key=f"zoom_{os.path.basename(file_path)}",
        )

        try:
            if selected_zoom_key == "Fit":
                st.image(file_path, use_container_width=True)
            elif selected_zoom_key == "100%":
                st.image(file_path, use_container_width=False)
            else:
                st.image(file_path, width=1200)
        except Exception as e:
            st.error(f"Error: {e}")

    def display_audio(self, file_path: str, format: str = "audio/mp3") -> None:
        """Ses dosyalarını arayüzde oynatır."""
        with open(file_path, "rb") as f:
            st.audio(f.read(), format=format)

    def display_video(self, file_path: str, format: str = "video/mp4") -> None:
        """Video dosyalarını arayüzde oynatır."""
        with open(file_path, "rb") as f:
            st.video(f.read(), format=format)

    def display_text_document(self, file_path: str, texts: dict | None = None):
        """TXT, DOCX ve kod/veri metin dosyalarını okuyup arayüzde gösterir."""
        if texts is None:
            texts = {}
            
        _, uzanti = os.path.splitext(file_path)
        uzanti = uzanti.lower()

        if uzanti in ['.docx', '.doc']:
            try:
                doc = docx.Document(file_path)
                tam_metin = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                st.markdown(tam_metin)
            except Exception as e:
                st.error(f"Error: {e}")

        elif uzanti == '.txt':
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                icerik = f.read()
            st.text_area(texts.get("document_content_label", "Content"), icerik, height=400)

        else:
            _lang_map = {
                ".py": "python", ".js": "javascript", ".html": "html",
                ".css": "css", ".java": "java", ".cpp": "cpp",
                ".sql": "sql", ".yaml": "yaml", ".json": "json",
                ".xml": "xml",
            }
            lang = _lang_map.get(uzanti, "text")
            try:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    icerik = f.read()
                st.code(icerik, language=lang)
            except Exception as e:
                st.error(f"Error: {e}")
